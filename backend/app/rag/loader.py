"""知识库加载：多格式解析 + 结构感知切块（中文优先分隔符 + 标题树）"""

import re
from pathlib import Path
from typing import Any, Callable

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config import config

# 分类目录 → category 标签
CATEGORY_DIRS = {
    "interview": "interview_questions",
    "project": "projects",
    "resume": "resume",
    "jd": "jd",
}
CATEGORY_NAMES = {
    "interview": "面试题库",
    "project": "项目文档",
    "resume": "个人简历",
    "jd": "岗位JD",
}

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_CHINESE_SEPARATORS = ["\n\n", "\n", "。", "；", "！", "？", "，", " ", ""]
# Q&A 型题目边界：数字/中文序号、Q+编号、第X题、【题目】等
_QA_BOUNDARY_RE = re.compile(
    r"^(?:(?:\d{1,3}|[一二三四五六七八九十]+)\s*[.、．)）]"
    r"|Q\s*\d+\s*[.、:：]?"
    r"|第\s*\d+\s*题"
    r"|【题目?\s*\d*\s*】)\s*(.*)$"
)


def _parse_qa_blocks(content: str) -> list[dict] | None:
    """Q&A 型切块：按题目边界整题切分，保留题号作 section。
    返回 None 表示非 Q&A 型（交给笔记型解析）。"""
    blocks: list[dict] = []
    cur_q: str | None = None
    cur_lines: list[str] = []
    boundary_count = 0
    for line in content.splitlines():
        s = line.strip()
        m = _QA_BOUNDARY_RE.match(s)
        if m:
            boundary_count += 1
            if cur_q:
                blocks.append({"content": "\n".join(cur_lines), "section": cur_q})
            cur_q = s[:60]
            cur_lines = [s]
        elif cur_q is not None:
            cur_lines.append(line.rstrip())
    if cur_q:
        blocks.append({"content": "\n".join(cur_lines), "section": cur_q})
    return blocks if boundary_count >= 2 else None


def _make_splitter() -> RecursiveCharacterTextSplitter:
    return RecursiveCharacterTextSplitter(
        chunk_size=config.kb_chunk_size,
        chunk_overlap=config.kb_chunk_overlap,
        separators=_CHINESE_SEPARATORS,
    )


def _split_long_block(content: str, section: str) -> list[dict[str, Any]]:
    if len(content) <= config.kb_chunk_size * 1.2:
        return [{"content": content, "section": section}]
    out = []
    for piece in _make_splitter().split_text(content):
        out.append({"content": piece, "section": section})
    return out


def _parse_markdown_blocks(content: str) -> list[dict[str, Any]]:
    """标题树切块：每个块带章节路径（如 手册 > 电机过热排查）"""
    blocks: list[dict[str, Any]] = []
    stack: list[str] = []
    cur: list[str] = []
    cur_section = ""

    def flush() -> None:
        nonlocal cur
        if cur:
            text = "\n".join(cur).strip()
            if text:
                blocks.append({"content": text, "section": cur_section})
            cur = []

    for line in content.splitlines():
        m = _HEADING_RE.match(line.rstrip())
        if m:
            flush()
            level = len(m.group(1))
            title = m.group(2).strip()
            stack = stack[: level - 1] + [title]
            cur_section = " > ".join(stack)
            cur.append(line.strip())
        else:
            cur.append(line.rstrip())
    flush()

    out: list[dict[str, Any]] = []
    for b in blocks:
        out.extend(_split_long_block(b["content"], b["section"]))
    return out


def _parse_text_blocks(content: str) -> list[dict[str, Any]]:
    """纯文本：按段落聚合 + 中文感知切分"""
    paragraphs = [p.strip() for p in content.splitlines() if p.strip()]
    if not paragraphs:
        return []
    blocks: list[dict[str, Any]] = []
    buf = ""
    for p in paragraphs:
        if buf and len(buf) + len(p) > config.kb_chunk_size:
            blocks.extend(_split_long_block(buf, "正文"))
            buf = ""
        buf = f"{buf}\n{p}" if buf else p
    if buf:
        blocks.extend(_split_long_block(buf, "正文"))
    return blocks


def _parse_html_blocks(content: str) -> list[dict[str, Any]]:
    """HTML：去标签还原标题结构后复用 Markdown 解析"""
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(content, "html.parser")
    lines: list[str] = []
    for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li"]):
        text = tag.get_text(" ", strip=True)
        if not text:
            continue
        if tag.name.startswith("h"):
            lines.append("#" * int(tag.name[1]) + " " + text)
        else:
            lines.append(text)
    return _parse_markdown_blocks("\n".join(lines))


def _parse_pdf_blocks(content: bytes) -> list[dict[str, Any]]:
    """PDF：pypdf 提正文（带页码）"""
    from pypdf import PdfReader

    import io

    blocks: list[dict[str, Any]] = []
    reader = PdfReader(io.BytesIO(content))
    for pno, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            blocks.extend(_split_long_block(text, f"第{pno}页"))
    return blocks


def _parse_docx_blocks(content: bytes) -> list[dict[str, Any]]:
    """DOCX：段落识别标题层级"""
    import io

    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(content))
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        m = re.match(r"(heading|标题)\s*(\d)", style)
        if m:
            lines.append("#" * min(int(m.group(2)), 6) + " " + text)
        else:
            lines.append(text)
    return _parse_markdown_blocks("\n".join(lines))


SUFFIX_HANDLERS: dict[str, Callable] = {
    ".md": _parse_markdown_blocks,
    ".markdown": _parse_markdown_blocks,
    ".txt": _parse_text_blocks,
    ".html": _parse_html_blocks,
    ".htm": _parse_html_blocks,
    ".pdf": _parse_pdf_blocks,
    ".docx": _parse_docx_blocks,
}


def infer_category(path: Path) -> str | None:
    """根据父目录名推断分类"""
    for cat, dirname in CATEGORY_DIRS.items():
        if dirname in path.parts:
            return cat
    return None


def parse_file(path: Path, category: str | None = None) -> list[Document]:
    """解析单个文件为带元数据的 Document 列表"""
    category = category or infer_category(path)
    if category is None:
        category = "project"
    suffix = path.suffix.lower()
    if suffix not in SUFFIX_HANDLERS:
        return []
    data: str | bytes
    if suffix in (".pdf", ".docx"):
        data = path.read_bytes()
    else:
        data = path.read_text(encoding="utf-8", errors="replace")
    blocks: list[dict] = []
    doc_type = "笔记型"
    # Q&A 型检测（md/txt 优先按题目边界切块，保留题号语义）
    if suffix in (".md", ".markdown", ".txt"):
        qa_blocks = _parse_qa_blocks(data)
        if qa_blocks:
            blocks, doc_type = qa_blocks, "Q&A"
    if not blocks:
        blocks = SUFFIX_HANDLERS[suffix](data)
    docs = []
    for i, b in enumerate(blocks):
        meta = {
            "source": path.name,
            "category": category,
            "category_name": CATEGORY_NAMES.get(category, category),
            "chunk_index": i,
            "doc_type": doc_type,
        }
        if b.get("section"):
            meta["section"] = b["section"]
        docs.append(Document(page_content=b["content"], metadata=meta))
    return docs


def load_data_dir() -> list[Document]:
    """扫描 data/ 下所有分类目录，返回全部 Document"""
    docs: list[Document] = []
    root = config.kb_data_path
    if not root.exists():
        return docs
    for cat, dirname in CATEGORY_DIRS.items():
        d = root / dirname
        if not d.exists():
            continue
        for f in sorted(d.rglob("*")):
            if f.is_file() and f.suffix.lower() in SUFFIX_HANDLERS:
                docs.extend(parse_file(f, cat))
    return docs
